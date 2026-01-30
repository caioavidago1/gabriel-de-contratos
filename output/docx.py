"""
Geração de relatório de análise DOCX.
Formato: 1) Problema  2) Parágrafo original  3) Parágrafo corrigido  4) Explicação das mudanças (quando houver).
"""
from docx import Document
from typing import List, Dict
from pathlib import Path
from datetime import datetime
import io



def gerar_relatorio_analise_docx(violacoes: List[Dict], gerar_sugestoes_reescrita: bool = False) -> bytes:
    """
    Gera um relatório de análise DOCX:
    1. Problema
    2. Parágrafo original
    3. Parágrafo corrigido (um abaixo do outro, sem diff)
    4. Explicação das mudanças (quando houver sugestão)


    Args:
        violacoes: Lista de violações com chunk, clausula_violada, motivo, sugestao_reescrita
        gerar_sugestoes_reescrita: Flag indicando se sugestões foram solicitadas


    Returns:
        Bytes do documento DOCX gerado
    """
    doc = Document()


    # Título principal
    doc.add_heading('Relatório de Análise de Contrato', 0)


    # Informações gerais
    doc.add_paragraph(f"Total de violações identificadas: {len(violacoes)}")
    doc.add_paragraph("")


    for i, violacao in enumerate(violacoes, 1):
        regra = violacao.get('regra', {})
        chunk = violacao.get('chunk', {})
        clausula_violada = regra.get('titulo', '') or violacao.get('clausula_violada', 'Regra não identificada')
        localizacao = violacao.get('localizacao') or (chunk.get('titulo') if chunk else None) or violacao.get('titulo', f'Violacao {i}')
        titulo_violacao = localizacao if localizacao and localizacao != 'N/A' else clausula_violada
        doc.add_heading(f'Violacao {i}: {titulo_violacao}', level=1)


        # 1. Problema
        doc.add_heading('1. Problema', level=2)


        motivos = violacao.get('motivo', [])
        problema_texto = violacao.get('problema', '').strip()
        if not problema_texto and motivos:
            problema_texto = "\n".join(f"• {m}" for m in motivos if m) if isinstance(motivos, list) else str(motivos)
        if not problema_texto:
            analise_agent1 = violacao.get('analise_agent1', '')
            if analise_agent1:
                problema_texto = analise_agent1.replace("[IA]", "").strip()
                problema_texto = problema_texto.replace("CLASSIFICAÇÃO: VIOLAÇÃO", "").strip()
                problema_texto = problema_texto.replace("CLASSIFICAÇÃO:VIOLAÇÃO", "").strip()
                problema_texto = problema_texto.replace("CLASSIFICAÇÃO: VIOLACAO", "").strip()
        if not problema_texto:
            problema_texto = f"Violacao da regra: {clausula_violada}"


        para_problema = doc.add_paragraph(problema_texto)
        para_problema.style = 'Normal'
        doc.add_paragraph("")


        # 2. Parágrafo original / 3. Parágrafo corrigido (um abaixo do outro, sem diff)
        chunk = violacao.get('chunk', {})
        texto_original = chunk.get('texto', '') if chunk else ''
        sugestao = violacao.get('sugestao_reescrita', {})
        texto_reescrito = sugestao.get('texto_reescrito', '') if sugestao and isinstance(sugestao, dict) else ''


        if not texto_reescrito and not gerar_sugestoes_reescrita:
            texto_reescrito = (
                "Sugestão de reescrita não solicitada. "
                "Ative a opção 'Gerar sugestões de reescrita' e execute nova análise para obter propostas de redação."
            )
        elif not texto_reescrito and gerar_sugestoes_reescrita:
            texto_reescrito = (
                "Falha na geração da sugestão de reescrita para esta violação. Verifique os logs ou tente novamente."
            )


        doc.add_heading('2. Parágrafo original', level=2)
        para_orig = doc.add_paragraph(texto_original or "(vazio)")
        para_orig.style = 'Normal'


        doc.add_paragraph("")
        doc.add_heading('3. Parágrafo corrigido', level=2)
        para_corrigido = doc.add_paragraph(texto_reescrito or "(vazio)")
        para_corrigido.style = 'Normal'


        # 4. Explicação das mudanças (quando houver sugestão)
        if sugestao and isinstance(sugestao, dict):
            explicacao = sugestao.get('explicacao_mudancas', '')
            if explicacao:
                doc.add_paragraph("")
                doc.add_heading('4. Explicação das mudanças', level=2)
                para_explicacao = doc.add_paragraph(explicacao)
                para_explicacao.style = 'Normal'


        if i < len(violacoes):
            doc.add_paragraph("")
            doc.add_paragraph("─" * 50)
            doc.add_paragraph("")


    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



def gerar_nome_relatorio_analise(nome_original: str) -> str:
    """Gera nome do arquivo de relatório de análise baseado no original."""
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_base = Path(nome_original).stem
    return f"relatorio_analise_{nome_base}_{timestamp}.docx"



def _formato_clausula(i: int) -> str:
    """Formato 'Cláusula N -' para problemas e solução."""
    return f"Cláusula {i} -"


def _adicionar_separador_fixo(doc: Document, numero_clausula: int, total_clausulas: int):
    """
    Adiciona separador fixo que permanece igual nos dois documentos.
    Isso garante que o track changes reconheça a estrutura correspondente.
    
    Args:
        doc: Documento onde adicionar o separador
        numero_clausula: Número da cláusula atual
        total_clausulas: Total de cláusulas no documento
    """
    if numero_clausula < total_clausulas:
        # Adiciona linha divisória em itálico (mesmo texto em ambos os docs)
        separador = doc.add_paragraph()
        run = separador.add_run(f"--- Fim da Cláusula {numero_clausula} ---")
        run.italic = True
        run.font.size = None  # Usa tamanho padrão
        
        # Adiciona linha em branco
        doc.add_paragraph("")


def _adicionar_cabecalho_clausula(doc: Document, numero_clausula: int):
    """
    Adiciona cabeçalho fixo da cláusula que será igual em ambos os documentos.
    
    Args:
        doc: Documento onde adicionar o cabeçalho
        numero_clausula: Número da cláusula
    """
    # Identificador único e fixo para cada cláusula
    cabecalho = doc.add_paragraph()
    run = cabecalho.add_run(f"[CLÁUSULA {numero_clausula:03d}]")
    run.bold = True
    run.font.size = None


def gerar_problemas_docx(violacoes: List[Dict]) -> bytes:
    """
    Gera DOCX apenas com o texto das cláusulas problemáticas.
    Inclui contexto fixo entre cláusulas para garantir alinhamento no track changes.
    
    Estrutura:
    [CLÁUSULA 001]
    Cláusula 1 -
    texto original
    --- Fim da Cláusula 1 ---
    
    [CLÁUSULA 002]
    Cláusula 2 -
    texto original
    --- Fim da Cláusula 2 ---
    """
    doc = Document()
    total = len(violacoes)
    
    for i, v in enumerate(violacoes, 1):
        # 1. Adiciona cabeçalho fixo da cláusula
        _adicionar_cabecalho_clausula(doc, i)
        
        # 2. Adiciona identificador "Cláusula N -"
        doc.add_paragraph(_formato_clausula(i))
        
        # 3. Adiciona o texto original
        chunk = v.get("chunk", {})
        texto = chunk.get("texto", "") if chunk else ""
        doc.add_paragraph(texto)
        
        # 4. Adiciona separador fixo (não será modificado no doc de solução)
        _adicionar_separador_fixo(doc, i, total)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



def gerar_solucao_docx(violacoes: List[Dict]) -> bytes:
    """
    Gera DOCX com o texto das cláusulas reescritas/corrigidas.
    Mantém EXATAMENTE a mesma estrutura que gerar_problemas_docx para garantir
    que o track changes funcione corretamente.
    
    Estrutura (IDÊNTICA ao doc de problemas):
    [CLÁUSULA 001]
    Cláusula 1 -
    texto revisado
    --- Fim da Cláusula 1 ---
    
    [CLÁUSULA 002]
    Cláusula 2 -
    texto revisado
    --- Fim da Cláusula 2 ---
    """
    doc = Document()
    total = len(violacoes)
    
    for i, v in enumerate(violacoes, 1):
        # 1. Adiciona cabeçalho fixo da cláusula (MESMO do doc de problemas)
        _adicionar_cabecalho_clausula(doc, i)
        
        # 2. Adiciona identificador "Cláusula N -" (MESMO do doc de problemas)
        doc.add_paragraph(_formato_clausula(i))
        
        # 3. Adiciona o texto revisado (ÚNICA diferença entre os documentos)
        sugestao = v.get("sugestao_reescrita", {})
        texto = sugestao.get("texto_reescrito", "") if isinstance(sugestao, dict) else ""
        if not texto:
            chunk = v.get("chunk", {})
            texto = chunk.get("texto", "") if chunk else "(Reescrita não gerada.)"
        doc.add_paragraph(texto)
        
        # 4. Adiciona separador fixo (MESMO do doc de problemas)
        _adicionar_separador_fixo(doc, i, total)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
