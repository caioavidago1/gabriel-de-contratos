"""
Comparar dois DOCX (problemas x solução) com track changes.
Usa win32com no Windows e LibreOffice UNO no Linux.
"""
import io
import os
import sys
import platform


if platform.system() != 'Windows':
    # Adiciona o caminho do LibreOffice ao sys.path para encontrar 'uno'
    libreoffice_path = '/usr/lib/python3/dist-packages'
    if libreoffice_path not in sys.path:
        sys.path.insert(0, libreoffice_path)
    
    # ✅ IMPORTS UNO AQUI (FORA DAS FUNÇÕES)
    import uno
    from com.sun.star.beans import PropertyValue
    from com.sun.star.connection import NoConnectException


from typing import Optional
from docx import Document


def gerar_doc_comparado(doc_problemas_bytes: bytes, doc_solucao_bytes: bytes) -> Optional[bytes]:
    """
    Gera um DOCX que concatena problemas + solução (usado pelo orquestrador para "explicacao").
    Não usa Word; retorna bytes do DOCX.
    """
    try:
        doc_problemas = Document(io.BytesIO(doc_problemas_bytes))
        doc_solucao = Document(io.BytesIO(doc_solucao_bytes))
    except Exception:
        return None
    doc = Document()
    doc.add_heading("Comparação: Problemas x Solução", 0)
    doc.add_paragraph("")
    doc.add_heading("1. Documento Problemas (original)", level=1)
    for para in doc_problemas.paragraphs:
        p = doc.add_paragraph(para.text)
        p.style = para.style.name if para.style else "Normal"
    doc.add_paragraph("")
    doc.add_heading("2. Documento Solução (revisado)", level=1)
    for para in doc_solucao.paragraphs:
        p = doc.add_paragraph(para.text)
        p.style = para.style.name if para.style else "Normal"
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def comparar_docx_windows(original_path, revisado_path, saida_path):
    """Comparação usando win32com (Windows/MS Word)"""
    import pythoncom
    import win32com.client as win32
    
    original_path = os.path.abspath(original_path)
    revisado_path = os.path.abspath(revisado_path)
    saida_path = os.path.abspath(saida_path)

    # Necessário quando chamado de outro thread (ex.: Streamlit)
    pythoncom.CoInitialize()
    try:
        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False

        try:
            doc_original = word.Documents.Open(original_path)
            doc_revisado = word.Documents.Open(revisado_path)

            # Equivalente ao Word: Revisar -> Comparar
            word.CompareDocuments(
                doc_original,
                doc_revisado,
                Destination=win32.constants.wdCompareDestinationNew,
                Granularity=win32.constants.wdGranularityWordLevel,
                CompareFormatting=True
            )

            # O resultado fica como ActiveDocument (documento de comparação)
            word.ActiveDocument.SaveAs(saida_path)

        finally:
            # Fecha tudo sem salvar os originais
            try:
                for d in list(word.Documents):
                    d.Close(SaveChanges=False)
            except Exception:
                pass
            word.Quit()
    finally:
        pythoncom.CoUninitialize()


def comparar_docx_linux(original_path, revisado_path, saida_path):
    """Comparação usando LibreOffice UNO (Linux)"""
    import uno
    from com.sun.star.beans import PropertyValue
    from com.sun.star.connection import NoConnectException
    
    def criar_propriedade(nome, valor):
        """Cria uma PropertyValue para passar argumentos UNO"""
        prop = PropertyValue()
        prop.Name = nome
        prop.Value = valor
        return prop
    
    original_path = os.path.abspath(original_path)
    revisado_path = os.path.abspath(revisado_path)
    saida_path = os.path.abspath(saida_path)
    
    try:
        # Conecta ao LibreOffice
        local_context = uno.getComponentContext()
        resolver = local_context.ServiceManager.createInstanceWithContext(
            "com.sun.star.bridge.UnoUrlResolver", local_context)
        
        ctx = resolver.resolve(
            "uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext")
        smgr = ctx.ServiceManager
        
        # Cria o desktop
        desktop = smgr.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
        
        # Converte caminhos para URLs do sistema
        original_url = uno.systemPathToFileUrl(original_path)
        revisado_url = uno.systemPathToFileUrl(revisado_path)
        saida_url = uno.systemPathToFileUrl(saida_path)
        
        # Abre o documento mais recente (revisor)
        propriedades = (criar_propriedade("Hidden", True),)
        documento = desktop.loadComponentFromURL(revisado_url, "_blank", 0, propriedades)
        
        # Obtém o dispatcher
        dispatcher = smgr.createInstanceWithContext(
            "com.sun.star.frame.DispatchHelper", ctx)
        
        # Prepara os argumentos para comparação
        args = (
            criar_propriedade("URL", original_url),
            criar_propriedade("FilterName", ""),
        )
        
        # Executa o comando de comparação
        frame = documento.getCurrentController().getFrame()
        dispatcher.executeDispatch(frame, ".uno:CompareDocuments", "", 0, args)
        
        # Salva o resultado
        props_salvar = (
            criar_propriedade("FilterName", "Office Open XML Text"),
            criar_propriedade("Overwrite", True),
        )
        documento.storeToURL(saida_url, props_salvar)
        
        # Fecha o documento
        documento.close(True)
        
    except NoConnectException:
        raise RuntimeError(
            "Erro: LibreOffice não está rodando. Execute:\n"
            'soffice --invisible --headless "--accept=socket,host=localhost,port=2002;urp;"&'
        )
    except Exception as e:
        raise RuntimeError(f"Erro ao comparar documentos: {e}")


def comparar_docx(original_path, revisado_path, saida_path):
    """
    Comparar dois DOCX com track changes (multiplataforma).
    Detecta automaticamente o sistema operacional.
    """
    if platform.system() == 'Windows':
        comparar_docx_windows(original_path, revisado_path, saida_path)
    else:
        comparar_docx_linux(original_path, revisado_path, saida_path)


if __name__ == "__main__":
    # Diretório padrão: output/docs (ao lado deste script, pasta docs)
    docs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")

    if len(sys.argv) >= 4:
        problema = sys.argv[1]
        solucao = sys.argv[2]
        saida = sys.argv[3]
        
        # Se forem só nomes de arquivo, buscar em output/docs
        if not os.path.isabs(problema) and not os.path.dirname(problema):
            problema = os.path.join(docs_dir, problema)
        if not os.path.isabs(solucao) and not os.path.dirname(solucao):
            solucao = os.path.join(docs_dir, solucao)
        if not os.path.isabs(saida) and not os.path.dirname(saida):
            saida = os.path.join(docs_dir, saida)
        
        try:
            comparar_docx(problema, solucao, saida)
            print(f"✓ Comparação concluída! Resultado salvo em: {saida}")
        except Exception as e:
            print(f"✗ Erro: {e}", file=sys.stderr)
            sys.exit(1)
    else:
        print("Uso: python comparar_docx.py <problemas_xxx.docx> <solucao_xxx.docx> <saida.docx>")
        print("Ex.: python comparar_docx.py problemas_Teste.docx solucao_Teste.docx comparacao.docx")
